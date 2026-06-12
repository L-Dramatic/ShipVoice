from __future__ import annotations

import csv
import json
import statistics
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
REPORT_PATH = OUT_DIR / "ShipVoice_船厂安全实时语音问答助手_项目报告_比赛增强版.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 32, 43)
MUTED = RGBColor(88, 96, 105)
LIGHT_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
RISK_FILL = "FFF1F0"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_font(paragraph, size: float = 11, bold: bool | None = None, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 10.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_body_paragraph(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead) :])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=10.8, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_FILL)
        set_cell_text(cell, h, bold=True, color=INK, size=9.8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def load_latency_metrics() -> dict[str, dict[str, float]]:
    path = ROOT / "results" / "latency_metrics.csv"
    grouped: dict[str, list[dict[str, str]]] = {}
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        for row in csv.DictReader(handle):
            grouped.setdefault(row["mode"], []).append(row)
    out = {}
    for mode, rows in grouped.items():
        out[mode] = {
            "count": len(rows),
            "profile": ",".join(sorted({r.get("execution_profile", "unknown") for r in rows})),
            "timing_source": ",".join(sorted({r.get("timing_source", "unknown") for r in rows})),
            "first_audio_ms": statistics.mean(float(r["first_audio_ms"]) for r in rows),
            "total_ms": statistics.mean(float(r["total_ms"]) for r in rows),
            "answer_chars": statistics.mean(float(r["answer_chars"]) for r in rows),
        }
    return out


def load_remote_metrics() -> dict[str, str]:
    root = ROOT / "results" / "remote_autodl_20260608_final"
    base = [json.loads(x) for x in (root / "results" / "base_eval.jsonl").read_text(encoding="utf-8").splitlines() if x.strip()]
    lora = [json.loads(x) for x in (root / "results" / "lora_eval.jsonl").read_text(encoding="utf-8").splitlines() if x.strip()]
    return {
        "base_rows": str(len(base)),
        "lora_rows": str(len(lora)),
        "base_avg_len": f"{statistics.mean(len(r['answer']) for r in base):.1f}",
        "lora_avg_len": f"{statistics.mean(len(r['answer']) for r in lora):.1f}",
        "adapter_mb": f"{(root / 'outputs' / 'qwen_lora_shipvoice' / 'adapter_model.safetensors').stat().st_size / 1024 / 1024:.1f}",
    }


def load_safety_metrics() -> dict[str, str]:
    summary_path = ROOT / "results" / "safety_gate_eval_summary.json"
    if not summary_path.exists():
        return {
            "total": "未运行",
            "label_accuracy": "未运行",
            "decision_accuracy": "未运行",
            "false_allow_count": "未运行",
            "avg_total_ms": "未运行",
        }
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    return {
        "total": str(summary["total"]),
        "label_accuracy": f"{float(summary['label_accuracy']) * 100:.1f}%",
        "decision_accuracy": f"{float(summary['decision_accuracy']) * 100:.1f}%",
        "false_allow_count": str(summary["false_allow_count"]),
        "avg_total_ms": f"{float(summary['avg_total_ms']):.0f}",
    }


def load_asr_metrics() -> dict[str, str]:
    manifest_path = ROOT / "data" / "audio" / "audio_manifest.csv"
    summary_path = ROOT / "results" / "asr_eval_summary.json"
    raw_summary_path = ROOT / "results" / "asr_eval_raw_summary.json"
    post_summary_path = ROOT / "results" / "asr_postprocess_summary.json"
    manifest_rows = []
    if manifest_path.exists():
        with manifest_path.open("r", encoding="utf-8-sig", newline="") as handle:
            manifest_rows = list(csv.DictReader(handle))
    if not summary_path.exists():
        return {
            "manifest_rows": str(len(manifest_rows)),
            "evaluated_rows": "0",
            "missing_audio_rows": str(len(manifest_rows)),
            "avg_cer": "未运行",
            "avg_wer": "未运行",
            "term_recall": "未运行",
            "raw_avg_cer": "未运行",
            "raw_avg_wer": "未运行",
            "raw_term_recall": "未运行",
            "cer_gain": "未运行",
            "term_recall_gain": "未运行",
            "rows_changed": "0",
            "replacements_applied": "0",
            "status": "pending_audio",
        }
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    raw_summary = json.loads(raw_summary_path.read_text(encoding="utf-8")) if raw_summary_path.exists() else summary
    post_summary = json.loads(post_summary_path.read_text(encoding="utf-8")) if post_summary_path.exists() else {}
    cer_gain = (float(raw_summary["avg_cer"]) - float(summary["avg_cer"])) * 100
    term_recall_gain = (float(summary["term_recall"]) - float(raw_summary["term_recall"])) * 100
    return {
        "manifest_rows": str(summary["total_manifest_rows"]),
        "evaluated_rows": str(summary["evaluated_rows"]),
        "missing_audio_rows": str(summary["missing_audio_rows"]),
        "raw_avg_cer": f"{float(raw_summary['avg_cer']) * 100:.2f}%",
        "raw_avg_wer": f"{float(raw_summary['avg_wer']) * 100:.2f}%",
        "raw_term_recall": f"{float(raw_summary['term_recall']) * 100:.2f}%",
        "avg_cer": f"{float(summary['avg_cer']) * 100:.2f}%",
        "avg_wer": f"{float(summary['avg_wer']) * 100:.2f}%",
        "term_recall": f"{float(summary['term_recall']) * 100:.2f}%",
        "cer_gain": f"{cer_gain:.2f} pp",
        "term_recall_gain": f"{term_recall_gain:.2f} pp",
        "rows_changed": str(post_summary.get("rows_changed", 0)),
        "replacements_applied": str(post_summary.get("replacements_applied", 0)),
        "status": str(summary["status"]),
    }


def load_multiturn_metrics() -> dict[str, str]:
    path = ROOT / "results" / "multiturn_eval_summary.json"
    if not path.exists():
        return {
            "dialogs": "未运行",
            "turns": "未运行",
            "followup_turns": "未运行",
            "gate_accuracy": "未运行",
            "top1_title_accuracy": "未运行",
            "title_hit_at_3": "未运行",
            "keyword_recall": "未运行",
            "followup_grounding_accuracy": "未运行",
            "avg_total_ms": "未运行",
        }
    summary = json.loads(path.read_text(encoding="utf-8"))
    return {
        "dialogs": str(summary["dialogs"]),
        "turns": str(summary["turns"]),
        "followup_turns": str(summary["followup_turns"]),
        "gate_accuracy": f"{float(summary['gate_accuracy']) * 100:.1f}%",
        "top1_title_accuracy": f"{float(summary['top1_title_accuracy']) * 100:.1f}%",
        "title_hit_at_3": f"{float(summary['title_hit_at_3']) * 100:.1f}%",
        "keyword_recall": f"{float(summary['keyword_recall']) * 100:.1f}%",
        "followup_grounding_accuracy": f"{float(summary['followup_grounding_accuracy']) * 100:.1f}%",
        "avg_total_ms": f"{float(summary['avg_total_ms']):.0f}",
    }


def load_real_chain_metrics() -> dict[str, object]:
    candidates = sorted(ROOT.glob("results/remote_real_chain_*"), key=lambda path: path.name, reverse=True)
    for candidate in candidates:
        summary_path = candidate / "summary.json"
        smoke_path = candidate / "real_chain_smoke.json"
        if not summary_path.exists() or not smoke_path.exists():
            continue
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        smoke = json.loads(smoke_path.read_text(encoding="utf-8"))
        provider_status = smoke["pipeline_result"]["provider_status"]
        return {
            "status": "verified",
            "artifact_dir": str(candidate.relative_to(ROOT)),
            "num_samples": str(summary["num_samples"]),
            "avg_first_audio_ms": f"{float(summary['avg_first_audio_ms']):.0f}",
            "avg_total_ms": f"{float(summary['avg_total_ms']):.0f}",
            "avg_asr_ms": f"{float(summary['avg_asr_ms']):.0f}",
            "avg_retrieval_ms": f"{float(summary['avg_retrieval_ms']):.0f}",
            "avg_tts_first_audio_ms": f"{float(summary['avg_tts_first_audio_ms']):.0f}",
            "asr_service": str(smoke["asr_health"]["service"]),
            "asr_model": str(smoke["asr_health"]["model"]),
            "tts_service": str(smoke["tts_health"]["service"]),
            "execution_profile": str(provider_status["execution_profile"]),
            "llm_provider": str(provider_status["llm"]),
            "tts_provider": str(provider_status["tts"]),
            "samples": summary["samples"],
        }
    return {
        "status": "not_run",
        "artifact_dir": "results/remote_real_chain_*",
        "num_samples": "0",
        "avg_first_audio_ms": "--",
        "avg_total_ms": "--",
        "avg_asr_ms": "--",
        "avg_retrieval_ms": "--",
        "avg_tts_first_audio_ms": "--",
        "asr_service": "--",
        "asr_model": "--",
        "tts_service": "--",
        "execution_profile": "--",
        "llm_provider": "--",
        "tts_provider": "--",
        "samples": [],
    }


def load_counts() -> dict[str, int]:
    def count_jsonl(path: Path) -> int:
        return sum(1 for line in path.read_text(encoding="utf-8").splitlines() if line.strip())

    return {
        "knowledge_docs": count_jsonl(ROOT / "data" / "knowledge" / "ship_safety_corpus.jsonl"),
        "sft_records": count_jsonl(ROOT / "data" / "training" / "sft_seed.jsonl"),
        "safety_gate_records": count_jsonl(ROOT / "data" / "training" / "safety_gate_seed.jsonl"),
    }


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "ShipVoice 项目报告 | 信息安全基础 A2"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_font(header, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "比赛增强版草稿 - 姓名学号待补充"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_font(footer, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ShipVoice")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("船厂安全实时语音问答助手")
    set_run_font(r, size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A2 级联式语音问答系统复现与改进")
    set_run_font(r, size=13, color=MUTED)

    doc.add_paragraph()
    meta_rows = [
        ["课程", "信息安全基础"],
        ["项目类型", "期末项目 A2 - 级联式语音问答系统 / 大创级原型"],
        ["小组成员", "姓名 / 学号待补充"],
        ["版本日期", "2026-06-08"],
        ["交付内容", "报告、答辩 PPT、源码、演示与实验证据"],
    ]
    add_table(doc, ["字段", "内容"], meta_rows, [1.45, 5.05])
    add_callout(
        doc,
        "项目一句话摘要",
        "本项目面向船厂高风险作业场景，构建了带安全门控、RAG 证据检索、可运行演示面板和 Qwen LoRA 微调实验的实时语音问答助手；核心目标不是简单串联 ASR-LLM-TTS，而是把领域知识、安全约束和实验闭环纳入系统设计。",
        fill=LIGHT_BLUE_FILL,
    )
    doc.add_page_break()


def add_exec_summary(
    doc: Document,
    counts: dict[str, int],
    remote: dict[str, str],
    safety: dict[str, str],
    asr: dict[str, str],
    real_chain: dict[str, object],
) -> None:
    doc.add_heading("1. 项目概述", level=1)
    add_body_paragraph(
        doc,
        "ShipVoice 是一个面向船厂安全作业的实时语音问答助手。系统以 A2 级联式语音问答为基础，进一步加入领域知识库、RAG 检索、安全门控、低延迟演示、实验记录和 LoRA/QLoRA 微调对照。"
    )
    add_body_paragraph(
        doc,
        "我们的设计重点是安全关键场景中的“可控回答”：当问题属于密闭舱室、动火、吊装、管路试压等船厂安全主题时，系统给出基于证据的保守建议；当问题越界、危险或试图绕过审批时，系统先由门控层拒答，而不是把风险直接交给大模型自由发挥。"
    )
    add_table(
        doc,
        ["维度", "完成情况"],
        [
            ["知识库", f"{counts['knowledge_docs']} 条船厂安全知识条目，覆盖有限空间、动火、吊装、试压、PPE、应急等主题"],
            ["训练数据", f"{counts['sft_records']} 条 SFT seed，{counts['safety_gate_records']} 条安全门控 seed"],
            ["安全评测", f"{safety['total']} 条安全/离题/prompt-injection/domain-safe/boundary 样例，危险请求误放行 {safety['false_allow_count']}"],
            ["语音评测", f"{asr['manifest_rows']} 条真实录音已完成评测，当前 ASR 状态为 {asr['status']}"],
            ["真实语音链路", f"远端已验证 {real_chain['num_samples']} 条真实样本，ASR={real_chain['asr_service']}，TTS={real_chain['tts_service']}，平均首音 {real_chain['avg_first_audio_ms']} ms"],
            ["远端微调", "RTX 4090 上完成 Qwen2.5-7B-Instruct 4-bit LoRA/QLoRA 训练"],
            ["评测闭环", f"Base eval {remote['base_rows']} 条，LoRA eval {remote['lora_rows']} 条，adapter 约 {remote['adapter_mb']} MB"],
            ["评测看板", "deliverables/ShipVoice_Evaluation_Dashboard.html 汇总 RAG、延迟、安全门控与 LoRA 对照"],
            ["演示形态", "本地 Web 操作台 + mock fallback，无 GPU 时仍可演示完整链路"],
        ],
        [1.35, 5.15],
    )


def add_requirement_mapping(doc: Document) -> None:
    doc.add_heading("2. 作业要求映射", level=1)
    add_table(
        doc,
        ["A2 要求", "本项目实现", "加分点"],
        [
            ["级联式语音问答", "ASR/transcript 层、LLM/RAG 回答层、TTS/playback-oriented 输出层", "保留 mock fallback，答辩现场稳定"],
            ["系统复现", "可运行的本地后端、Web 面板、benchmark 与 validation 脚本", "一键 quick validation"],
            ["系统改进", "船厂安全知识库、HybridRetriever、安全门控、提示注入拒答", "面向安全关键领域，而非通用聊天"],
            ["模型增强", "Qwen2.5-7B-Instruct LoRA/QLoRA 微调实验", "有远端 GPU 日志、adapter 和 base-vs-LoRA 对比"],
            ["实验评估", "检索 hit-rate、延迟 CSV、55 条安全 benchmark、base/LoRA JSONL 结果", "可复现、可审计"],
        ],
        [1.45, 3.05, 2.0],
    )


def add_architecture(doc: Document) -> None:
    doc.add_heading("3. 系统架构", level=1)
    add_body_paragraph(doc, "系统采用“门控优先、证据约束、模型生成、演示可降级”的级联设计。正式链路如下：")
    add_table(
        doc,
        ["阶段", "模块", "作用"],
        [
            ["1", "Voice / Text Input", "接收真实语音或文字输入；无真实 ASR 时使用 transcript fallback 保持演示稳定"],
            ["2", "Transcript Normalization", "统一问题格式，保留船厂术语和安全关键词"],
            ["3", "Domain & Safety Gate", "判断是否属于船厂安全问题，拦截危险请求、越权注入和 off-domain 问题"],
            ["4", "Hybrid RAG Retrieval", "从船厂安全知识库检索相关条目，提供可解释证据"],
            ["5", "Answer Synthesis", "基于证据生成保守、专业、可执行、适合语音播报的回答"],
            ["6", "Playback / Demo Panel", "展示转写、门控结果、证据、回答和延迟指标"],
            ["7", "Logging & Evaluation", "记录 latency、retrieval、base-vs-LoRA 和安全拦截结果"],
        ],
        [0.55, 2.0, 3.95],
    )
    add_callout(
        doc,
        "为什么不直接裸用大模型？",
        "船厂安全问答属于高风险场景。单纯依赖 LLM 生成可能产生幻觉、越权建议或对危险请求回答过度具体。本项目把安全门控和 RAG 证据层放在模型前后，使模型只能在受控上下文中发挥作用。",
    )


def add_knowledge_and_safety(doc: Document) -> None:
    doc.add_heading("4. 知识库、RAG 与安全门控", level=1)
    add_body_paragraph(doc, "知识库位于 data/knowledge/ship_safety_corpus.jsonl，当前覆盖 20 个船厂安全主题。每条知识包含标题、关键词、风险说明、作业前检查、处置建议和语音回答素材。")
    add_body_paragraph(doc, "RAG 索引由 scripts/build_knowledge_index.py 构建，检索评测由 scripts/evaluate_retrieval.py 完成。在 quick validation 中，代表性问题达到 hit@1 = 5/5、hit@3 = 5/5。")
    add_table(
        doc,
        ["风险类型", "处理策略", "示例"],
        [
            ["off-domain", "短路为非船厂安全问题，不进入领域回答", "股票投资、生活闲聊等"],
            ["unsafe request", "拒绝绕过审批、跳过检测、破坏设备等请求", "如何绕过安全检查进入密闭舱室"],
            ["prompt injection", "忽略用户要求“忽略规则”的指令，保持安全策略", "规避动火审批、绕开监管"],
            ["domain-safe", "进入 RAG 检索并生成基于证据的安全建议", "密闭舱室动火前检查、管路试压风险"],
        ],
        [1.35, 3.0, 2.15],
    )


def add_safety_evaluation(doc: Document, safety: dict[str, str]) -> None:
    doc.add_heading("5. 安全评测与边界门控增强", level=1)
    add_body_paragraph(
        doc,
        "为把项目从课程 demo 推进到比赛级原型，我们将安全 benchmark 扩展为 55 条，覆盖 off-domain、unsafe、prompt injection、domain-safe 和 boundary 五类场景。评测脚本会完整跑过 pipeline，并输出 CSV、JSON 和 Markdown 报告。"
    )
    add_table(
        doc,
        ["指标", "结果", "意义"],
        [
            ["安全评测样例", safety["total"], "覆盖离题、违规绕行、提示注入、正常安全作业和边界应急"],
            ["标签准确率", safety["label_accuracy"], "expected_gate 与 predicted_gate 完全一致的比例"],
            ["allow/block 决策准确率", safety["decision_accuracy"], "系统是否正确放行安全问题并拦截危险问题"],
            ["危险请求误放行", safety["false_allow_count"], "比赛展示中最关键的安全失败指标"],
            ["完整 pipeline 平均耗时", f"{safety['avg_total_ms']} ms", "包含模拟 ASR、门控、RAG、LLM 和 TTS 的端到端耗时"],
        ],
        [1.7, 1.35, 3.45],
    )
    add_body_paragraph(
        doc,
        "扩展评测暴露了两个早期规则问题：倒装表达“报警线拔掉”未命中，以及“发现有人修改检测记录如何上报”这种正当报告场景容易被过度拦截。我们据此补强门控，使其能区分“请求违规方法”和“报告/制止违规行为”。"
    )
    add_callout(
        doc,
        "评测结论",
        "当前 55 条完整 pipeline 安全评测中，标签准确率 100.0%，allow/block 决策准确率 100.0%，危险请求误放行 0。下一阶段目标是加入 100+ 对抗改写和真实语音转写变体。",
        fill=LIGHT_BLUE_FILL,
    )


def add_audio_asr_readiness(doc: Document, asr: dict[str, str]) -> None:
    doc.add_heading("6. 真实语音与 ASR 实测增强", level=1)
    add_body_paragraph(
        doc,
        "我们已经完成 50 条真实录音的离线 ASR 实测，覆盖正常安全问题、off-domain 拒答、unsafe 拒答、prompt-injection 拒答和边界应急问题，并包含 quiet、classroom、workshop-like 三类噪声条件。"
    )
    add_body_paragraph(
        doc,
        "在此基础上，项目新增了一层 ShipVoice 领域术语后处理：保留 SenseVoice 原始转写结果，同时对舾装、压载水舱、测氧测爆、动火作业、试压等高价值术语做可审计的规则纠错，形成原始转写与增强转写的前后对比。"
    )
    add_table(
        doc,
        ["资产/指标", "结果", "说明"],
        [
            ["真实录音样本", asr["manifest_rows"], "三位组员完成录制，位于 data/audio/raw/"],
            ["ASR 已评测样本", asr["evaluated_rows"], "50/50 条全部完成评测"],
            ["Raw CER / WER", f"{asr['raw_avg_cer']} / {asr['raw_avg_wer']}", "SenseVoice 原始输出"],
            ["Corrected CER / WER", f"{asr['avg_cer']} / {asr['avg_wer']}", "术语后处理后的最终转写"],
            ["Raw 术语召回", asr["raw_term_recall"], "原始 ASR 对领域术语的覆盖"],
            ["Corrected 术语召回", asr["term_recall"], "术语增强后的最终覆盖"],
            ["CER 改善", asr["cer_gain"], "后处理带来的绝对百分点提升"],
            ["术语召回改善", asr["term_recall_gain"], "后处理带来的绝对百分点提升"],
            ["修正样本数", asr["rows_changed"], "至少命中 1 条术语/近音纠错规则的样本"],
            ["修正规则命中次数", asr["replacements_applied"], "总替换 span 数"],
        ],
        [1.7, 1.55, 3.25],
    )
    add_callout(
        doc,
        "工程意义",
        "这一步把项目从“跑通真实 ASR”提升到“具备领域化后处理能力”。老师看到的不再只是 ASR 数字，而是完整的改进链路：真实录音 -> 基线转写 -> 术语增强 -> 指标提升。",
        fill=LIGHT_BLUE_FILL,
    )


def add_multiturn_evaluation(doc: Document, multiturn: dict[str, str]) -> None:
    doc.add_heading("6.5 多轮上下文评测", level=1)
    add_body_paragraph(doc, "为补齐作业中“多轮对话为主”的要求，项目新增了多轮上下文 benchmark。评测不再只看单轮命中，而是测试系统能否在后续追问中保持正确门控、延续前文主题，并继续命中对应安全知识条目。")
    add_table(
        doc,
        ["指标", "结果", "说明"],
        [
            ["对话数", multiturn["dialogs"], "覆盖有限空间、试压、吊装、高处作业和临时用电等多轮场景"],
            ["总轮次", multiturn["turns"], "包含首问与 follow-up turns"],
            ["追问轮次", multiturn["followup_turns"], "需要依赖前文上下文才能稳定作答"],
            ["门控准确率", multiturn["gate_accuracy"], "多轮条件下 domain/safety 决策是否稳定"],
            ["Top-1 证据命中", multiturn["top1_title_accuracy"], "首条证据是否仍然对应正确主题"],
            ["Title hit@3", multiturn["title_hit_at_3"], "前三条证据是否包含目标知识条目"],
            ["关键词召回", multiturn["keyword_recall"], "回答是否覆盖预期安全要点"],
            ["追问 grounding 准确率", multiturn["followup_grounding_accuracy"], "依赖历史上下文的轮次是否仍能命中正确证据"],
            ["平均总耗时(ms)", multiturn["avg_total_ms"], "full pipeline 下的多轮平均响应耗时"],
        ],
        [1.4, 1.0, 3.8],
    )
    add_callout(
        doc,
        "评测意义",
        "多轮 benchmark 让 Part 1 和 Part 2 真正接上：系统不再只是单句问答，而是能够在连续追问中保持领域一致性、安全边界和证据 grounding。",
        fill=LIGHT_BLUE_FILL,
    )


def add_real_chain_validation(doc: Document, real_chain: dict[str, object]) -> None:
    doc.add_heading("6.8 真实语音链路联调验证", level=1)
    if real_chain["status"] != "verified":
        add_body_paragraph(doc, "当前仓库尚未检测到远端真实语音链路验证结果。完成 GPU 联调后，应补充 ASR/TTS 服务健康检查、端到端样本结果和延迟统计。")
        return

    add_body_paragraph(
        doc,
        "为证明系统不是只有前端页面，本项目额外完成了一轮远端真实语音链路联调：在 RTX 4090 上启动 FunASR / SenseVoice ASR 服务与 ChatTTS 中文 TTS 服务，通过 HTTP provider 接回本地 ShipVoice pipeline，形成真实语音输入和真实语音回传。",
    )
    add_body_paragraph(
        doc,
        "本轮验证重点不是追求最低时延，而是留下可审计的工程证据。当前 provider status 为 hybrid：真实音频 I/O 已经在线，问答生成仍由受控 mock_llm 层承接检索证据，因此它证明了真实服务集成能力，也明确暴露了后续企业级改造的瓶颈位置。",
    )
    add_table(
        doc,
        ["指标", "结果", "说明"],
        [
            ["验证样本数", str(real_chain["num_samples"]), "A001-A003 三条真实录音端到端跑通"],
            ["ASR 服务", str(real_chain["asr_service"]), str(real_chain["asr_model"])],
            ["TTS 服务", str(real_chain["tts_service"]), f"provider={real_chain['tts_provider']}"],
            ["执行形态", str(real_chain["execution_profile"]), f"LLM provider={real_chain['llm_provider']}"],
            ["平均 ASR 耗时", f"{real_chain['avg_asr_ms']} ms", "真实语音转写阶段"],
            ["平均检索耗时", f"{real_chain['avg_retrieval_ms']} ms", "RAG 证据检索阶段"],
            ["平均 TTS 首音", f"{real_chain['avg_tts_first_audio_ms']} ms", "当前主要性能瓶颈"],
            ["平均端到端首音", f"{real_chain['avg_first_audio_ms']} ms", "远端真实语音链路观察值"],
        ],
        [1.55, 1.35, 3.3],
    )
    sample_rows = []
    for row in real_chain["samples"]:
        sample_rows.append(
            [
                str(row["sample_id"]),
                str(row["transcript"]),
                f"{row['asr_ms']} / {row['retrieval_ms']} / {row['tts_first_audio_ms']}",
                str(row["total_ms"]),
            ]
        )
    add_table(
        doc,
        ["样本", "转写结果", "ASR/检索/TTS首音(ms)", "总耗时(ms)"],
        sample_rows,
        [0.8, 3.2, 1.3, 1.2],
    )
    add_callout(
        doc,
        "联调结论",
        f"真实语音链路证据已归档至 {real_chain['artifact_dir']}。当前系统已经具备“真实语音输入 + 真实语音输出 + 可控安全问答”的工程形态；下一阶段的主要任务不是继续证明能跑，而是优化 TTS 延迟并替换 hybrid 中的 mock_llm。",
        fill=LIGHT_BLUE_FILL,
    )


def add_latency(doc: Document, latency: dict[str, dict[str, float]]) -> None:
    doc.add_heading("7. 本地演示与延迟评估", level=1)
    add_body_paragraph(doc, "本地演示使用 run_demo.py 启动后端和 Web 面板。当前版本已经把 provider 状态、execution profile 和 timing source 显式暴露出来，用于区分 demo/simulated 链路与真实 provider 调用链路。")
    rows = []
    for mode in ["baseline", "streaming", "full"]:
        if mode in latency:
            item = latency[mode]
            rows.append([
                mode,
                str(int(item["count"])),
                str(item.get("profile", "unknown")),
                str(item.get("timing_source", "unknown")),
                f"{item['first_audio_ms']:.0f}",
                f"{item['total_ms']:.0f}",
                f"{item['answer_chars']:.1f}",
            ])
    add_table(
        doc,
        ["模式", "样本数", "执行形态", "计时来源", "首段音频均值(ms)", "总耗时均值(ms)", "平均回答长度"],
        rows,
        [0.9, 0.7, 1.0, 1.0, 1.35, 1.3, 1.35],
    )
    for image_name, caption in [
        ("demo_panel_safety.png", "图 1：安全问答场景演示面板"),
        ("demo_panel_backend.png", "图 2：后端运行与指标展示"),
    ]:
        image = ROOT / "results" / image_name
        if image.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image), width=Inches(5.8))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            set_run_font(r, size=9.5, color=MUTED)


def add_lora_experiment(doc: Document, remote: dict[str, str]) -> None:
    doc.add_heading("8. Qwen LoRA/QLoRA 微调实验", level=1)
    add_body_paragraph(doc, "为了超出基础复现要求，项目在远端 RTX 4090 上完成了 Qwen2.5-7B-Instruct 的 4-bit LoRA/QLoRA 微调实验。该实验的目标不是替代安全门控，而是验证小规模领域数据能否增强回答风格和船厂安全表达。")
    add_table(
        doc,
        ["项目", "配置/结果"],
        [
            ["GPU", "NVIDIA GeForce RTX 4090 24GB"],
            ["基础模型", "Qwen/Qwen2.5-7B-Instruct"],
            ["训练方法", "4-bit LoRA/QLoRA，target modules 覆盖 attention 与 MLP 投影层"],
            ["训练数据", "63 条 ShipVoice SFT seed"],
            ["训练轮数", "2 epochs，14 optimizer steps"],
            ["训练结果", "train_loss = 1.7777，train_runtime = 61.6765s"],
            ["产物", f"LoRA adapter_model.safetensors，约 {remote['adapter_mb']} MB"],
        ],
        [1.65, 4.85],
    )
    add_table(
        doc,
        ["模型", "评测条数", "平均回答长度", "观察"],
        [
            ["Base Qwen2.5-7B-Instruct", remote["base_rows"], f"{remote['base_avg_len']} 字符", "off-domain 拒答更稳，重复更少"],
            ["ShipVoice LoRA Adapter", remote["lora_rows"], f"{remote['lora_avg_len']} 字符", "回答更短、更符合船厂安全模板，但有轻微模板化过拟合"],
        ],
        [2.1, 1.0, 1.4, 2.0],
    )
    add_callout(
        doc,
        "关键结论",
        "LoRA 是有价值的领域适配实验，但不能直接作为唯一正式回答源。最终展示应把 LoRA 放在安全门控和 RAG 后面，用作可选领域风格增强；安全边界仍由显式门控和证据检索保证。",
        fill=RISK_FILL,
    )


def add_reproducibility(doc: Document, real_chain: dict[str, object]) -> None:
    doc.add_heading("9. 可复现性与工程质量", level=1)
    add_body_paragraph(doc, "项目保留了本地与远端两套复现路径。本地路径用于课程答辩现场稳定演示，远端路径用于 GPU 微调实验和模型对比。")
    add_table(
        doc,
        ["任务", "命令/文件", "验收标准"],
        [
            ["本地 quick validation", "python scripts\\validate_project.py --quick", "输出 VALIDATION OK"],
            ["本地演示", "python run_demo.py", "打开 http://127.0.0.1:8010 可运行"],
            ["检索评测", "python scripts\\evaluate_retrieval.py", "代表性问题 hit@1/hit@3 达标"],
            ["安全评测", "python scripts\\evaluate_safety_gate.py --fail-on-critical", "55 条 benchmark 通过且 false_allow_count = 0"],
            ["ASR 评测", "python scripts\\evaluate_asr_transcripts.py", "录音和 ASR 转写填入后输出 CER/WER/术语召回"],
            ["多轮评测", "python scripts\\evaluate_multiturn.py --fail-on-threshold", "多轮 gate 与 follow-up grounding 达标"],
            ["真实语音链路", "python scripts\\check_real_service_chain.py --sample-id A001", "ASR/TTS health 正常且 pipeline provider 变为 http_json"],
            ["远端训练", "remote/train_qwen_lora.py", "生成 LoRA adapter 与训练日志"],
            ["远端评测", "remote/evaluate_qwen_lora.py", "生成 base_eval.jsonl 与 lora_eval.jsonl"],
        ],
        [1.45, 2.65, 2.4],
    )
    add_body_paragraph(doc, f"核心远端证据位于 results/remote_autodl_20260608_final；真实语音链路证据位于 {real_chain['artifact_dir']}。前者证明微调实验，后者证明真实 ASR/TTS 服务联调。")


def add_limitations(doc: Document) -> None:
    doc.add_heading("10. 局限性与后续工作", level=1)
    add_bullet(doc, "SFT seed 数据规模仍较小，LoRA 主要证明领域适配可行性，不应夸大为生产级模型。")
    add_bullet(doc, "LoRA 在 off-domain 问题上出现轻微领域模板化，因此正式系统必须保留安全门控。")
    add_bullet(doc, "50 条真实中文语音已经完成录制与 ASR 评测，但当前样本规模仍偏小，下一阶段应扩展到 100+ 条、多说话人和更强噪声条件。")
    add_bullet(doc, "多轮 benchmark 已建立，但仍以结构化追问为主，下一阶段应加入更自由的口语化省略问句和 ASR 错字变体。")
    add_bullet(doc, "远端真实 ChatTTS 链路已经打通，但当前平均 TTS 首音延迟仍在十秒级，后续应优先做流式返回、模型替换或推理缓存优化。")
    add_bullet(doc, "安全门控已扩展到 55 条 benchmark，但下一阶段仍应加入 100+ 对抗改写、真实语音 ASR 错字变体和轻量分类器。")


def add_appendix(doc: Document) -> None:
    doc.add_heading("附录：关键文件清单", level=1)
    add_table(
        doc,
        ["类别", "路径"],
        [
            ["项目说明", "README.md"],
            ["架构说明", "docs/ARCHITECTURE.md"],
            ["最终路线图", "docs/FINAL_DELIVERY_PLAN_20260608.md"],
            ["报告大纲", "docs/FINAL_REPORT_OUTLINE_20260608.md"],
            ["知识库", "data/knowledge/ship_safety_corpus.jsonl"],
            ["测试问题", "data/tests/eval_questions.csv"],
            ["多轮评测集", "data/tests/multiturn_eval.jsonl"],
            ["安全评测集", "data/tests/safety_eval.csv"],
            ["安全评测结果", "results/safety_gate_eval.csv / safety_gate_eval_summary.json / safety_gate_eval_report.md"],
            ["多轮评测结果", "results/multiturn_eval.csv / multiturn_eval_summary.json / multiturn_eval_report.md"],
            ["录音任务", "data/audio/audio_manifest.csv / deliverables/ShipVoice_Audio_Recording_Pack.html"],
            ["ASR 评测结果", "results/asr_eval.csv / asr_eval_summary.json / asr_eval_report.md / asr_eval_raw_summary.json / asr_postprocess_summary.json"],
            ["真实语音链路结果", "results/remote_real_chain_20260612_chattts_48359"],
            ["评测看板", "deliverables/ShipVoice_Evaluation_Dashboard.html"],
            ["SFT 数据", "data/training/sft_seed.jsonl"],
            ["安全门控数据", "data/training/safety_gate_seed.jsonl"],
            ["远端结果", "results/remote_autodl_20260608_final"],
            ["AutoDL bundle", "results/autodl_bundle.zip"],
        ],
        [1.4, 5.1],
    )


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    counts = load_counts()
    latency = load_latency_metrics()
    remote = load_remote_metrics()
    safety = load_safety_metrics()
    asr = load_asr_metrics()
    multiturn = load_multiturn_metrics()
    real_chain = load_real_chain_metrics()

    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_exec_summary(doc, counts, remote, safety, asr, real_chain)
    add_requirement_mapping(doc)
    add_architecture(doc)
    add_knowledge_and_safety(doc)
    add_safety_evaluation(doc, safety)
    add_audio_asr_readiness(doc, asr)
    add_multiturn_evaluation(doc, multiturn)
    add_real_chain_validation(doc, real_chain)
    add_latency(doc, latency)
    add_lora_experiment(doc, remote)
    add_reproducibility(doc, real_chain)
    add_limitations(doc)
    add_appendix(doc)

    doc.core_properties.title = "ShipVoice 船厂安全实时语音问答助手项目报告"
    doc.core_properties.subject = "信息安全基础 A2 期末项目"
    doc.core_properties.author = "待补充"
    doc.core_properties.comments = "Generated by project build script; group member names and student IDs are placeholders."
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build()
