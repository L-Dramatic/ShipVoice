"""Build the final ShipVoice course report DOCX from the curated Markdown file.

This script is intentionally separate from the runtime application. It exists
only to regenerate the final submission document with stable Word formatting.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "deliverables" / "final_submission" / "report"
REPORT_MD = REPORT_DIR / "ShipVoice_船厂安全实时语音问答助手_项目报告_最终版.md"
REPORT_DOCX = REPORT_MD.with_suffix(".docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)
HEADER_FILL = "F4F6F9"
CALLOUT_FILL = "F7F9FC"
CODE_FILL = "F2F4F7"


def ensure_rfonts_from_rpr(rpr):
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    return rfonts


def set_run_rfonts(run, name="Calibri", east="Microsoft YaHei"):
    rpr = run._r.get_or_add_rPr()
    rfonts = ensure_rfonts_from_rpr(rpr)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east)


def set_style_rfonts(style, name="Calibri", east="Microsoft YaHei"):
    rpr = style._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style._element.append(rpr)
    rfonts = ensure_rfonts_from_rpr(rpr)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    set_run_rfonts(run, name=name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def paragraph_bottom_border(paragraph, color="D9E2F3", size="4"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_rich_text(paragraph, text, size=None, bold_default=False):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5 if size is None else size, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold_default)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_width(table, widths):
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), "9360")
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def widths_for_cols(count):
    if count == 2:
        return [2800, 6560]
    if count == 3:
        return [2200, 2600, 4560]
    if count == 4:
        return [1700, 2200, 3100, 2360]
    if count == 5:
        return [1350, 1750, 2050, 2050, 2160]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, widths_for_cols(cols))
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_rich_text(paragraph, value.strip(), size=9.5, bold_default=True)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_rich_text(paragraph, value.strip(), size=9.3)
    doc.add_paragraph()


def parse_table_line(line):
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_separator(line):
    value = line.strip().strip("|").strip()
    return bool(value) and all(set(part.strip()) <= set("-: ") and "-" in part for part in value.split("|"))


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    set_style_rfonts(normal)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        set_style_rfonts(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "ShipVoice 船厂安全实时语音问答助手 | 信息安全基础期末项目 A2"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "ShipVoice 船厂安全实时语音问答助手 | 信息安全基础 A2"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)


def add_cover(doc, markdown_text):
    lines = [line.strip() for line in markdown_text.splitlines() if line.strip()]
    title = lines[0].lstrip("#").strip()
    metadata = []
    for line in lines[1:]:
        if line.startswith("## "):
            break
        if "：" in line:
            key, value = line.split("：", 1)
            metadata.append((key.strip(), value.strip()))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("ShipVoice")
    set_run_font(run, size=26, color=DARK_BLUE, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    set_run_font(run, size=21, color=BLACK, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = "信息安全基础期末项目 A2 | 级联式造船语音问答系统的复现与改进"
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    set_run_font(run, size=12, color=MUTED, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [2200, 7160])
    for row_index, (key, value) in enumerate(metadata):
        cells = table.rows[row_index].cells
        set_cell_shading(cells[0], HEADER_FILL)
        cells[0].text = ""
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich_text(p0, key, size=10.5, bold_default=True)
        cells[1].text = ""
        add_rich_text(cells[1].paragraphs[0], value, size=10.5)

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(18)
    callout.paragraph_format.space_after = Pt(8)
    shade_paragraph(callout, CALLOUT_FILL)
    add_rich_text(
        callout,
        "摘要：本报告围绕船厂安全语音问答场景，说明系统需求、架构设计、级联链路实现、安全控制、实验结果、复现方法和后续改进方向。",
        size=10.5,
    )
    paragraph_bottom_border(callout)
    doc.add_page_break()


def replace_mermaid_blocks(content: str) -> str:
    assets_dir = REPORT_DIR / "assets" / "diagrams"
    pattern = re.compile(r"```mermaid\n(.*?)```", re.DOTALL)

    def repl(match: re.Match[str]) -> str:
        for name in ("system_architecture.png", "system_architecture.svg"):
            image_path = assets_dir / name
            if image_path.exists():
                rel = Path("..") / "assets" / "diagrams" / name
                return f"![系统架构图]({rel.as_posix()})\n\n*图 3-1 ShipVoice 系统架构数据流*"
        return match.group(0)

    return pattern.sub(repl, content)


def add_markdown_image(doc, alt_text, image_path):
    resolved = (REPORT_DIR / image_path).resolve()
    if not resolved.exists():
        paragraph = doc.add_paragraph()
        add_rich_text(paragraph, f"[图片缺失] {alt_text} ({image_path})", size=10, bold_default=True)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(resolved), width=Inches(6.2))
    doc.add_paragraph()


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text.strip("* "))
    set_run_font(run, size=9.5, color=MUTED, italic=True)


def add_markdown_body(doc, markdown_text):
    markdown_text = replace_mermaid_blocks(markdown_text)
    lines = markdown_text.splitlines()
    start = 0
    for index, line in enumerate(lines):
        if line.startswith("## 摘要"):
            start = index
            break
    lines = lines[start:]

    in_code = False
    code_lines = []
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.05
            shade_paragraph(paragraph, CODE_FILL)
            run = paragraph.add_run("\n".join(code_lines))
            set_run_font(run, size=9, name="Consolas")
            code_lines = []

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if is_separator(line):
                continue
            table_rows.append(parse_table_line(line))
            continue
        flush_table()

        if not line.strip():
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line.strip())
        if image_match:
            add_markdown_image(doc, image_match.group(1), image_match.group(2))
            continue

        if line.strip().startswith("*图 ") and line.strip().endswith("*"):
            add_caption(doc, line.strip())
            continue

        if line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
        else:
            number_match = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
            if number_match:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.194)
                paragraph.paragraph_format.space_after = Pt(4)
                add_rich_text(paragraph, f"{number_match.group(1)}. {number_match.group(2)}")
            elif line.strip().startswith("- "):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.194)
                paragraph.paragraph_format.space_after = Pt(4)
                add_rich_text(paragraph, "• " + line.strip()[2:])
            else:
                paragraph = doc.add_paragraph()
                add_rich_text(paragraph, line.strip())

    flush_table()
    flush_code()


def main():
    markdown_text = REPORT_MD.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc, markdown_text)
    add_markdown_body(doc, markdown_text)
    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)
    print(REPORT_DOCX)


if __name__ == "__main__":
    main()
